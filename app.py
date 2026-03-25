from flask import Flask, render_template, jsonify, request, session
import json
import random
import os
import csv
from datetime import datetime
from docx import Document
import re

# Thử import AI Tutor
try:
    from ai_tutor import hoi_ai, cap_nhat_bo_nho_de
except ImportError:
    print("⚠️ Không tìm thấy file ai_tutor.py. Tính năng AI sẽ bị lỗi.")

app = Flask(__name__)
app.secret_key = "hoang-gia-dat-secret-key"

# Khởi tạo thư mục hệ thống
for path in ['static/uploads', 'data/quizzes', 'data/results', 'dataset']:
    os.makedirs(path, exist_ok=True)

# --- API HỆ THỐNG ---

# Mật khẩu giáo viên (Đạt có thể đổi tùy ý)
ADMIN_USER = "admin"
ADMIN_PASS = "admin123" 

# --- PHẦN ĐĂNG NHẬP GIÁO VIÊN ---

@app.route('/api/teacher-login', methods=['POST'])
def teacher_login_api():
    # Lấy dữ liệu JSON từ Fetch API gửi lên
    data = request.get_json()
    
    if not data:
        return jsonify({"status": "error", "message": "Không nhận được dữ liệu"}), 400

    # Lấy đúng tên biến 'user' và 'pass' từ JavaScript của bạn
    username = data.get('user')
    password = data.get('pass')

    print(f"--- Đạt ơi, Server nhận được: {username} / {password} ---")

    # Kiểm tra tài khoản (Đạt sửa admin/admin123 theo ý mình nhé)
    if username == "admin" and password == "admin123":
        session['is_teacher'] = True
        return jsonify({"status": "success"}), 200
    else:
        return jsonify({"status": "error", "message": "Sai tài khoản hoặc mật khẩu!"}), 401

# ROUTE DASHBOARD
@app.route('/teacher_dashboard')
def teacher_dashboard():
    if not session.get('is_teacher'):
        return redirect('/') # Nếu chưa login thì đá về trang chủ
    return render_template('teacher_dashboard.html')


# --- TRANG QUẢN LÝ (DASHBOARD) ---
@app.route('/teacher/dashboard') # Đảm bảo có dấu / ở giữa
def teacher_dashboard_final():
    if not session.get('is_teacher'):
        return redirect('/api/teacher-login')
    
    # Đọc file CSV để hiện điểm (Giữ nguyên logic của Đạt)
    return render_template('teacher_dashboard.html')

def parse_docx_to_questions(file_path):
    doc = Document(file_path)
    questions = []
    current_q = None
    
    # Regex nhận diện
    q_re = re.compile(r'^Câu\s*\d+[:.]', re.IGNORECASE)
    opt_re = re.compile(r'^[A-D]\.', re.IGNORECASE)
    tf_re = re.compile(r'^[a-d]\)', re.IGNORECASE)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text: continue

        if q_re.match(text):
            if current_q: questions.append(current_q)
            current_q = {
                "content": text, 
                "type": "multiple_choice", 
                "options": [], 
                "sub_questions": [], 
                "answer": "", # Để trống để cập nhật sau
                "explanation": ""
            }
            continue

        if current_q:
            if opt_re.match(text):
                current_q["options"].append(text)
                current_q["type"] = "multiple_choice"
            
            elif tf_re.match(text):
                current_q["type"] = "true_false"
                current_q["sub_questions"].append({"text": text, "answer": ""})

            elif "Đáp án:" in text or "ĐÁP ÁN:" in text:
                ans_val = text.split(":")[-1].strip()
                
                # NẾU LÀ ĐÚNG SAI: Tách chuỗi "a-Đ, b-S" thành mảng đáp án
                if current_q["type"] == "true_false":
                    current_q["answer"] = ans_val # Lưu chuỗi gốc để đối chiếu
                    # Cập nhật đáp án vào từng câu hỏi phụ
                    parts = ans_val.split(',')
                    for p in parts:
                        p = p.strip()
                        if '-' in p:
                            idx_char, val = p.split('-')
                            # Chuyển a, b, c, d thành 0, 1, 2, 3
                            idx = ord(idx_char.lower()) - ord('a')
                            if idx < len(current_q["sub_questions"]):
                                current_q["sub_questions"][idx]["answer"] = val.upper()
                
                # NẾU LÀ TRẮC NGHIỆM HOẶC TRẢ LỜI NGẮN
                else:
                    current_q["answer"] = ans_val
                    if not current_q["options"]:
                        current_q["type"] = "short_answer"

    if current_q: questions.append(current_q)
    return questions

@app.route('/api/admin/import-word', methods=['POST'])
def admin_import_word():
    try:
        if 'file' not in request.files:
            return jsonify({"status": "error", "message": "Không tìm thấy file"}), 400
        
        file = request.files['file']
        file_path = f"temp_{file.filename}"
        file.save(file_path)

        # GỌI HÀM ĐÃ ĐỊNH NGHĨA Ở BƯỚC 2
        questions = parse_docx_to_questions(file_path)

        # Xóa file tạm sau khi đọc xong
        if os.path.exists(file_path):
            os.remove(file_path)

        return jsonify({
            "status": "success",
            "questions": questions
        })
    except Exception as e:
        print(f"LỖI HỆ THỐNG: {e}")
        return jsonify({"status": "error", "message": str(e)}), 500

#def test_parse(file_path):
#    doc = Document(file_path)
#    for i, para in enumerate(doc.paragraphs):
#        text = para.text.strip()
#        if not text: continue
#        
#        # Kiểm tra nhận diện câu hỏi
#        if "Câu" in text:
#            print(f"--- Phát hiện Câu hỏi tại dòng {i}: {text[:30]}...")
#        
#        # Kiểm tra nhận diện đáp án
#        elif text.startswith(("A.", "B.", "C.", "D.")):
#            print(f"   [Lựa chọn]: {text}")
#            
#        # Kiểm tra nhận diện đáp án đúng
#        elif "Đáp án:" in text:
#            print(f"   >> ĐÁP ÁN ĐÚNG: {text.split(':')[-1].strip()}")

# Chạy thử với file của bạn
#test_parse('de_thi.docx')

@app.route('/teacher/dashboard')
def teacher_dashboard_view(): # Đổi tên hàm thành teacher_dashboard_view
    if not session.get('is_teacher'):
        return redirect('/teacher/login')
    return render_template('teacher_dashboard.html')

@app.route('/api/set-session', methods=['POST'])
def set_session():
    data = request.get_json()
    session['user_name'] = data.get('name')
    session['user_class'] = data.get('class')
    return jsonify({"status": "success"})

@app.route('/api/upload-word', methods=['POST'])
def upload_word():
    file = request.files.get('file')
    if not file: return jsonify({"error": "No file"}), 400
    doc = Document(file)
    questions = []
    current_q = None
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text: continue
        if text.lower().startswith("câu"):
            if current_q: questions.append(current_q)
            current_q = {"q": text, "options": [], "ans": "A"}
        elif current_q and any(text.upper().startswith(x) for x in ["A.", "B.", "C.", "D."]):
            current_q["options"].append(text)
    if current_q: questions.append(current_q)
    return jsonify({"questions": questions})


@app.route('/api/admin/save-quiz', methods=['POST'])
def admin_save_quiz(): 
    try:
        data = request.get_json()
        quiz_code = data.get('quiz_code', 'TEST').upper()

        # Thống nhất lưu vào data/quizzes giống hàm cũ của bạn
        os.makedirs("data/quizzes", exist_ok=True)

        file_path = f"data/quizzes/{quiz_code}.json"
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=4)

        print(f"--- Đã lưu đề thi Admin: {quiz_code} thành công! ---")
        return jsonify({"status": "success", "message": "Đã lưu đề thi!"})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route('/api/admin/list-quizzes')
def list_quizzes():
    path = "data/quizzes"
    if not os.path.exists(path):
        os.makedirs(path, exist_ok=True)
        return jsonify([])
    # Lấy tên các file .json trong thư mục
    files = [f.replace('.json', '') for f in os.listdir(path) if f.endswith('.json')]
    return jsonify(files)

@app.route('/api/admin/view-results')
def view_results():
    path = "data/results"
    if not os.path.exists(path):
        os.makedirs(path, exist_ok=True)
        return "<h3>Chưa có dữ liệu điểm thi.</h3><button onclick='history.back()'>Quay lại</button>"
    
    files = [f for f in os.listdir(path) if f.endswith('.csv')]
    if not files:
        return "<h3>Thư mục kết quả đang trống.</h3><button onclick='history.back()'>Quay lại</button>"

    # Tạo bảng hiển thị đơn giản
    html = "<h2>Danh sách bảng điểm học sinh</h2><ul>"
    for f in files:
        # Đường dẫn tải file trực tiếp
        html += f"<li style='margin:10px 0;'>{f} - <a href='/api/admin/download-result/{f}'>Tải file .CSV</a></li>"
    html += "</ul><br><button onclick='history.back()'>Quay lại Dashboard</button>"
    return html

    # API Lấy danh sách các file điểm CSV
@app.route('/api/admin/list-results')
def list_results():
    path = "data/results"
    if not os.path.exists(path):
        os.makedirs(path, exist_ok=True)
    
    # Lấy danh sách file .csv
    files = [f for f in os.listdir(path) if f.endswith('.csv')]
    return jsonify(files)

# API Đọc nội dung 1 file CSV cụ thể (để hiện bảng)
@app.route('/api/admin/get-result-detail')
def get_result_detail():
    filename = request.args.get('file')
    path = os.path.join("data/results", filename)
    
    if not os.path.exists(path):
        return jsonify({"error": "File không tồn tại"}), 404
        
    results = []
    import csv
    with open(path, mode='r', encoding='utf-8-sig') as f:
        reader = csv.reader(f)
        for row in reader:
            results.append(row) # Mỗi row là [Thời gian, Tên, Lớp, Điểm, Chi tiết]
            
    return jsonify(results)

@app.route('/api/submit-exam', methods=['POST'])
def api_submit_exam():
    try:
        data = request.json
        user_answers = data.get('answers')      # Mảng đáp án học sinh gửi lên
        quiz_code = data.get('quiz_code')        # Mã đề thi
        student_name = data.get('student_name', 'Ẩn danh')
        student_class = data.get('student_class', 'Không rõ')

        # 1. KIỂM TRA VÀ ĐỌC FILE ĐỀ THI (Fix lỗi NameError)
        file_path = f"data/quizzes/{quiz_code}.json"
        if not os.path.exists(file_path):
            return jsonify({"status": "error", "message": "Không tìm thấy mã đề này!"}), 404

        with open(file_path, 'r', encoding='utf-8') as f:
            quiz_data = json.load(f)

        score = 0.0
        details = []

        # 2. VÒNG LẶP CHẤM ĐIỂM 3 DẠNG ĐỀ
        for i, q in enumerate(quiz_data['questions']):
            # Lấy đáp án của học sinh cho câu này, nếu không có thì mặc định là chuỗi trống hoặc mảng trống
            u_ans = user_answers[i] if i < len(user_answers) else ""
            q_type = q.get('type', 'multiple_choice')

            # --- DẠNG 1: TRẮC NGHIỆM 4 LỰA CHỌN ---
            if q_type == 'multiple_choice':
                is_correct = (str(u_ans).strip().upper() == str(q.get('answer', '')).strip().upper())
                if is_correct: 
                    score += 0.25 # Mỗi câu trắc nghiệm thường 0.25đ
                details.append({"type": q_type, "is_correct": is_correct, "correct": q.get('answer')})

            # --- DẠNG 2: ĐÚNG SAI (Barem Bộ GD 2025) ---
            elif q_type == 'true_false':
                correct_count = 0
                sub_results = []
                # q['sub_questions'] chứa list các ý a, b, c, d
                for j, sub in enumerate(q.get('sub_questions', [])):
                    # u_ans lúc này phải là một mảng các lựa chọn [ "Đ", "S", "Đ", "S" ]
                    student_sub_ans = u_ans[j] if isinstance(u_ans, list) and j < len(u_ans) else ""
                    is_sub_correct = (str(student_sub_ans).strip().upper() == str(sub.get('answer', '')).strip().upper())
                    
                    if is_sub_correct: 
                        correct_count += 1
                    sub_results.append(is_sub_correct)
                
                # Barem điểm chuẩn: Đúng 1 ý: 0.1 | 2 ý: 0.25 | 3 ý: 0.5 | 4 ý: 1.0
                points = {1: 0.1, 2: 0.25, 3: 0.5, 4: 1.0}.get(correct_count, 0)
                score += points
                details.append({"type": q_type, "is_correct": (correct_count == 4), "correct_count": correct_count, "sub_results": sub_results})

            # --- DẠNG 3: TRẢ LỜI NGẮN ---
            elif q_type == 'short_answer':
                # Chuẩn hóa dấu phẩy thành dấu chấm để so sánh số thập phân
                clean_u = str(u_ans).strip().replace(',', '.')
                clean_q = str(q.get('answer', '')).strip().replace(',', '.')
                
                is_correct = (clean_u == clean_q)
                if is_correct: 
                    score += 0.5 # Câu trả lời ngắn thường điểm cao hơn
                details.append({"type": q_type, "is_correct": is_correct, "correct": q.get('answer')})

        # 3. GHI KẾT QUẢ VÀO FILE CSV ĐỂ QUẢN LÝ
        results_dir = "data/results"
        if not os.path.exists(results_dir):
            os.makedirs(results_dir)
            
        csv_file = f"{results_dir}/{quiz_code}_results.csv"
        file_exists = os.path.isfile(csv_file)
        
        with open(csv_file, 'a', newline='', encoding='utf-8-sig') as f:
            writer = csv.writer(f)
            # Viết tiêu đề nếu file mới
            if not file_exists:
                writer.writerow(["Thời gian", "Họ tên", "Lớp", "Điểm số"])
            
            now = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
            writer.writerow([now, student_name, student_class, round(score, 2)])

        # 4. TRẢ KẾT QUẢ VỀ CHO HỌC SINH
        return jsonify({
            "status": "success", 
            "score": round(score, 2), 
            "total_questions": len(quiz_data['questions']), # Tên biến này phải khớp với JS
            "details": details
        })

    except Exception as e:
        print(f"LỖI HỆ THỐNG KHI CHẤM ĐIỂM: {e}")
        return jsonify({"status": "error", "message": "Có lỗi xảy ra khi chấm điểm."}), 500
# 3. API Hỗ trợ tải file CSV
@app.route('/api/admin/download-result/<filename>')
def download_result(filename):
    return send_from_directory("data/results", filename)


@app.route('/api/save-quiz', methods=['POST'])
def save_quiz():
    data = request.get_json()
    code = data['info']['code'].upper()
    with open(f"data/quizzes/{code}.json", 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=4)
    return jsonify({"status": "success"})

@app.route('/api/get-quiz')
def get_quiz():
    code = request.args.get('code', '').upper()
    path = f"data/quizzes/{code}.json"
    if os.path.exists(path):
        with open(path, 'r', encoding='utf-8') as f:
            return jsonify(json.load(f))
    return jsonify({"error": "Mã đề không tồn tại"}), 404

@app.route('/api/check-session')
def check_session():
    if 'user_name' in session:
        return jsonify({
            "logged_in": True, 
            "name": session['user_name'], 
            "class": session['user_class']
        })
    return jsonify({"logged_in": False})

#--------------
#--------------
#-- practice --
de_hien_tai = []

goi_y_hoc = {
    "dao_ham": "Ôn quy tắc đạo hàm (x^n)' = n*x^(n-1)",
    "nguyen_ham": "Ôn nguyên hàm ∫x^n dx = x^(n+1)/(n+1) + C",
    "logarit": "Ôn quy tắc logarit log(ab)=log a + log b",
    "cap_so_cong": "Ôn công thức a_n = a1 + (n-1)d",
    "cap_so_nhan": "Ôn công thức a_n = a1*q^(n-1)",
    "xac_suat": "Ôn xác suất có điều kiện và Bayes",
    "Vật lý nhiệt": "Ôn lại nội năng và các nguyên lý nhiệt động lực học.",
    "Các đẳng quá trình": "Ghi nhớ các định luật Boyle, Charles và Gay-Lussac.",
    "Từ trường": "Xem lại quy tắc bàn tay trái và công thức lực Lorentz.",
    "Vật lý hạt nhân": "Ôn tập định luật bảo toàn số khối và chu kỳ bán rã.",
    "Điện xoay chiều": "Ghi nhớ công thức tính tổng trở Z và các hiện tượng cộng hưởng.",
    "Sóng điện từ": "Ôn lại tính chất của các loại sóng vô tuyến và cấu tạo máy thu phát.",
    "Sóng cơ": "Xem lại định nghĩa bước sóng và phương trình truyền sóng."
}

def tai_du_lieu(mon):
    try:
        # GIỮ NGUYÊN ĐƯỜNG DẪN CŨ CỦA ĐẠT
        filename = "dataset/math.json" if mon == "toan" else "dataset/physics.json"
        if not os.path.exists(filename):
            print(f"❌ Không tìm thấy file: {filename}")
            return []
        with open(filename, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        print("Lỗi dataset:", e)
        return []

@app.route("/concepts")
def concepts():
    mon = request.args.get("mon")
    if mon == "toan":
        concepts = ["hàm số và đồ thị", "nguyên hàm tích phân", "hình học không gian", "khối tròn xoay", "tọa độ không gian", "số mũ và logarit", "số phức", "tổ hợp và xác suất", "đạo hàm", "cấp số cộng", "cấp số nhân", "tiệm cận"]
    else:
        concepts = ["Khí lí tưởng", "Vật lý hạt nhân", "Từ trường", "Vật lí nhiệt"]
    return jsonify(concepts)

@app.route("/tao_de")
def tao_de():
    global de_hien_tai
    mon = request.args.get("mon")
    concept = request.args.get("concept")
    so_cau = int(request.args.get("so_cau", 10))
    
    data = tai_du_lieu(mon)
    if concept and concept != "all":
        data = [q for q in data if q.get("concept") == concept]
    
    if not data: 
        return jsonify({"de": [], "message": "Không tìm thấy câu hỏi!"})

    so_luong_lay = min(so_cau, len(data))
    selected_qs = random.sample(data, so_luong_lay)
    
    de_sau_khi_tron = []
    for q in selected_qs:
        q_copy = q.copy()
        shuffled_options = list(q_copy["options"])
        random.shuffle(shuffled_options)
        q_copy["options"] = shuffled_options
        de_sau_khi_tron.append(q_copy)
    
    random.shuffle(de_sau_khi_tron)
    de_hien_tai = de_sau_khi_tron

    # --- THÊM DÒNG NÀY ĐỂ AI LƯU ĐỀ VÀO BỘ NHỚ ---
    cap_nhat_bo_nho_de(de_hien_tai) 

    return jsonify({"de": de_hien_tai})

@app.route("/nop_bai", methods=["POST"])
def nop_bai():
    global de_hien_tai
    data = request.json
    student_name = data.get("student_name", "Ẩn danh")
    cau_tra_loi = data.get("cau_tra_loi", [])
    
    if not de_hien_tai:
        return jsonify({"error": "Chưa tạo đề"}), 400

    score, results, weak = 0, [], {}
    for q, ans in zip(de_hien_tai, cau_tra_loi):
        dung = (ans == q["answer"])
        if dung: score += 1
        else:
            c = q.get("concept", "unknown")
            weak[c] = weak.get(c, 0) + 1
            
        results.append({
            "cau_hoi": q["question"], 
            "tra_loi_cua_ban": ans, 
            "dap_an_dung": q["answer"], 
            "dung": dung
        })
    
    # Lưu lịch sử thi (giữ nguyên)
    file_path = "lich_su_thi.csv"
    file_exists = os.path.isfile(file_path)
    thoi_gian = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    
    try:
        with open(file_path, mode="a", newline="", encoding="utf-8-sig") as f:
            writer = csv.writer(f)
            if not file_exists:
                writer.writerow(["Thời gian", "Thí sinh", "Số câu đúng", "Tổng số câu", "Phần trăm"])
            writer.writerow([thoi_gian, student_name, score, len(de_hien_tai), f"{round(score/len(de_hien_tai)*100, 2)}%"])
    except Exception as e:
        print("Lỗi lưu file CSV:", e)

    goi_y = [goi_y_hoc[k] for k in weak if k in goi_y_hoc]
    
    return jsonify({
        "diem": score, 
        "phan_tram": round(score/len(de_hien_tai)*100, 2), 
        "ket_qua": results, 
        "goi_y": goi_y
    })

@app.route("/chat_ai", methods=["POST"])
def chat_ai():
    try:
        data = request.json
        msg = data.get("message", "")
        user_name = data.get("user_name", "Hoàng Gia Đạt")
        
        # ĐỊNH NGHĨA BIẾN Ở ĐÂY ĐỂ KHÔNG BỊ LỖI 'NOT DEFINED'
        system_prompt = f"""
Mày là "Gia sư Cheems" - Huyền thoại dạy kèm Toán và Lý tại THPT Việt Đức. Đang dạy bro {user_name}.
Phong cách: Lầy lội, Gen Z (bro, đỉnh nóc, kịch trần, khum, chê).

[QUY TẮC TỐI CAO]:
1. Nếu tin nhắn có '[WAIT_MODE]': TUYỆT ĐỐI CẤM GIẢI BÀI. Chỉ chào và liệt kê số câu sai.
2. Chỉ giải bài khi học sinh hỏi 'Giải câu X'. 
3. Dùng $ cho công thức cùng dòng, $$ cho công thức xuống dòng.
"""

        # 1. Kiểm tra Mật lệnh để ép AI vào chế độ liệt kê
        if "[WAIT_MODE]" in msg:
            # Gửi tin nhắn trực tiếp cho hàm gọi AI của Đạt
            # Mình giả sử hàm của Đạt là hoi_ai(nội_dung, tên, chế_độ)
            response_text = hoi_ai(msg, user_name=user_name, mode="tutor")
            
        else:
            # 2. Chat bình thường
            response_text = hoi_ai(msg, user_name=user_name, mode="tutor")

        # 3. Dọn dẹp LaTeX trước khi trả về (Ép AI dùng đúng chuẩn $)
        if response_text:
            response_text = response_text.replace("\\[", "$$").replace("\\]", "$$")
            response_text = response_text.replace("\\(", "$").replace("\\)", "$")

        return jsonify({"answer": response_text})

    except Exception as e:
        # In lỗi ra màn hình đen (Terminal) để Đạt soi cho dễ
        print(f"Lỗi rồi bro Đạt ơi: {e}")
        return jsonify({"answer": f"Lỗi rồi bro: {str(e)}"}), 500




# --- ĐIỀU HƯỚNG TRANG (Kiểm tra kỹ tên file .html) ---
@app.route('/')
def index(): return render_template('index.html')

@app.route('/student_hub')
def student_hub(): return render_template('student_hub.html')

@app.route('/practice')
def practice_page(): return render_template('practice.html') # Đổi tên hàm tránh trùng

@app.route('/exam_lobby')
def exam_lobby(): return render_template('exam_lobby.html')

@app.route('/exam_room')
def exam_room(): return render_template('exam_room.html')


if __name__ == "__main__":
    app.run(host='0.0.0.0', port=5000, debug=True)